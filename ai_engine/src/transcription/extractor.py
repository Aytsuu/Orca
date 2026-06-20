from __future__ import annotations

import asyncio
import inspect
import os
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from time import monotonic
from typing import Any

from src.config import get_settings
from src.exceptions import ConfigurationError

IMAGE_PROMPT = "Describe all visible content including any text."
AUDIO_PROMPT = "Transcribe the audio as plain text."
VIDEO_PROMPT = "Transcribe the video as plain text, including visible on-screen text when relevant."


@dataclass
class ExtractionResult:
    text: str
    method: str


class UnsupportedMimeType(Exception):
    pass


class VideoTooLong(Exception):
    pass


class TranscriptExtractor:
    def __init__(
        self,
        *,
        api_key: str | None = None,
        async_client: Any | None = None,
        model: str | None = None,
        video_max_duration_seconds: int | None = None,
        file_activation_timeout_seconds: int = 120,
        poll_interval_seconds: int = 2,
    ) -> None:
        settings = get_settings()
        self._api_key = api_key or settings.llm_api_key
        self._async_client = async_client
        self._model = model or settings.llm_fast_model
        self._video_max_duration_seconds = (
            video_max_duration_seconds or settings.video_max_duration_seconds
        )
        self._file_activation_timeout_seconds = file_activation_timeout_seconds
        self._poll_interval_seconds = poll_interval_seconds

    async def _get_async_client(self):
        if self._async_client is None:
            if not self._api_key:
                raise ConfigurationError(
                    "LLM_API_KEY is required for Gemini transcript extraction."
                )
            try:
                from google import genai
            except ImportError as exc:  # pragma: no cover - environment-specific
                raise ConfigurationError("google-genai is not installed.") from exc
            self._async_client = genai.Client(api_key=self._api_key).aio
        return self._async_client

    async def extract(self, file_bytes: bytes, mime_type: str) -> ExtractionResult:
        normalized_mime_type = str(mime_type or "").strip().lower()
        if normalized_mime_type == "application/pdf":
            return ExtractionResult(text=self._extract_pdf(file_bytes), method="pdf")
        if normalized_mime_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return ExtractionResult(text=self._extract_docx(file_bytes), method="docx")
        if normalized_mime_type in {"text/plain", "text/markdown", "text/csv"}:
            return ExtractionResult(text=self._extract_plaintext(file_bytes), method="plaintext")
        if normalized_mime_type.startswith("image/"):
            return ExtractionResult(
                text=await self._generate_inline_media_text(
                    file_bytes,
                    mime_type=normalized_mime_type,
                    prompt=IMAGE_PROMPT,
                ),
                method="gemini_vision",
            )
        if normalized_mime_type.startswith("audio/"):
            return ExtractionResult(
                text=await self._generate_file_media_text(
                    file_bytes,
                    mime_type=normalized_mime_type,
                    prompt=AUDIO_PROMPT,
                ),
                method="gemini_audio",
            )
        if normalized_mime_type.startswith("video/"):
            return ExtractionResult(
                text=await self._generate_video_text(file_bytes, normalized_mime_type),
                method="gemini_video",
            )
        raise UnsupportedMimeType(normalized_mime_type or "unknown")

    def _extract_pdf(self, file_bytes: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - environment-specific
            raise ConfigurationError("pypdf is not installed.") from exc

        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join((page.extract_text() or "").strip() for page in reader.pages)
        return text.strip()

    def _extract_docx(self, file_bytes: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - environment-specific
            raise ConfigurationError("python-docx is not installed.") from exc

        document = Document(BytesIO(file_bytes))
        parts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    def _extract_plaintext(self, file_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace").strip()

    async def _generate_inline_media_text(
        self,
        file_bytes: bytes,
        *,
        mime_type: str,
        prompt: str,
    ) -> str:
        try:
            from google.genai import types
        except ImportError as exc:  # pragma: no cover - environment-specific
            raise ConfigurationError("google-genai is not installed.") from exc

        client = await self._get_async_client()
        response = await client.models.generate_content(
            model=self._model,
            contents=[types.Part.from_bytes(data=file_bytes, mime_type=mime_type), prompt],
        )
        return self._response_text(response)

    async def _generate_video_text(self, file_bytes: bytes, mime_type: str) -> str:
        uploaded_file = await self._upload_media_file(file_bytes, mime_type)
        try:
            activated = await self._wait_for_file_activation(uploaded_file)
            duration_seconds = self._video_duration_seconds(activated)
            if duration_seconds is not None and duration_seconds > self._video_max_duration_seconds:
                raise VideoTooLong(
                    f"Video duration {duration_seconds:.2f}s exceeds the configured limit."
                )
            response = await self._generate_content_from_uploaded_file(activated, VIDEO_PROMPT)
            return self._response_text(response)
        finally:
            await self._delete_uploaded_file(uploaded_file)

    async def _generate_file_media_text(
        self,
        file_bytes: bytes,
        *,
        mime_type: str,
        prompt: str,
    ) -> str:
        uploaded_file = await self._upload_media_file(file_bytes, mime_type)
        try:
            activated = await self._wait_for_file_activation(uploaded_file)
            response = await self._generate_content_from_uploaded_file(activated, prompt)
            return self._response_text(response)
        finally:
            await self._delete_uploaded_file(uploaded_file)

    async def _upload_media_file(self, file_bytes: bytes, mime_type: str):
        client = await self._get_async_client()
        suffix = Path(mime_type.replace("/", ".")).suffix or ".bin"
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            response = client.files.upload(file=temp_path)
            if inspect.isawaitable(response):
                response = await response
            return response
        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    async def _wait_for_file_activation(self, uploaded_file):
        client = await self._get_async_client()
        current = uploaded_file
        deadline = monotonic() + self._file_activation_timeout_seconds
        while True:
            state = getattr(getattr(current, "state", None), "name", None) or str(
                getattr(current, "state", "") or ""
            )
            normalized_state = state.upper()
            if normalized_state == "ACTIVE":
                return current
            if normalized_state == "FAILED":
                raise RuntimeError("Gemini Files API failed to activate the uploaded file.")
            if monotonic() >= deadline:
                raise RuntimeError("Timed out while waiting for Gemini Files API activation.")
            response = client.files.get(name=current.name)
            current = await response if inspect.isawaitable(response) else response
            await asyncio.sleep(self._poll_interval_seconds)

    async def _generate_content_from_uploaded_file(self, uploaded_file, prompt: str):
        client = await self._get_async_client()
        return await client.models.generate_content(
            model=self._model,
            contents=[uploaded_file, prompt],
        )

    async def _delete_uploaded_file(self, uploaded_file) -> None:
        if not getattr(uploaded_file, "name", None):
            return
        client = await self._get_async_client()
        response = client.files.delete(name=uploaded_file.name)
        if inspect.isawaitable(response):
            await response

    def _video_duration_seconds(self, uploaded_file) -> float | None:
        metadata = getattr(uploaded_file, "video_metadata", None)
        if metadata is None and hasattr(uploaded_file, "videoMetadata"):
            metadata = uploaded_file.videoMetadata
        if metadata is None and isinstance(uploaded_file, dict):
            metadata = uploaded_file.get("videoMetadata") or uploaded_file.get("video_metadata")
        duration = getattr(metadata, "video_duration", None)
        if duration is None and hasattr(metadata, "videoDuration"):
            duration = metadata.videoDuration
        if duration is None and isinstance(metadata, dict):
            duration = metadata.get("videoDuration") or metadata.get("video_duration")
        if not duration:
            return None
        return self._parse_duration_seconds(str(duration))

    def _parse_duration_seconds(self, value: str) -> float:
        normalized = value.strip().lower()
        if not normalized.endswith("s"):
            raise ValueError(f"Unsupported duration value: {value}")
        return float(normalized[:-1])

    def _response_text(self, response: Any) -> str:
        text = getattr(response, "text", None)
        if isinstance(text, str):
            return text.strip()
        candidates = getattr(response, "candidates", None) or []
        for candidate in candidates:
            content = getattr(candidate, "content", None)
            parts = getattr(content, "parts", None) or []
            for part in parts:
                part_text = getattr(part, "text", None)
                if isinstance(part_text, str) and part_text.strip():
                    return part_text.strip()
        raise ValueError("Gemini did not return transcript text.")
